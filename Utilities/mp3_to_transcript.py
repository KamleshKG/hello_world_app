"""
MP3 to Transcript Converter
Uses OpenAI Whisper (local, free, no API key needed)
Model: medium (best for technical content)

Requirements:
    pip install openai-whisper
    pip install python-docx
    # Also install ffmpeg:
    # Windows: winget install ffmpeg  OR  choco install ffmpeg
    # Ubuntu:  sudo apt install ffmpeg

Usage:
    python mp3_to_transcript.py audio.mp3
    python mp3_to_transcript.py audio.mp3 --output my_transcript
    python mp3_to_transcript.py audio.mp3 --format txt
    python mp3_to_transcript.py audio.mp3 --format docx
    python mp3_to_transcript.py audio.mp3 --format both   (default)
"""

import argparse
import os
import sys
import time
from pathlib import Path


def check_dependencies():
    """Check if required packages are installed."""
    missing = []
    try:
        import whisper
    except ImportError:
        missing.append("openai-whisper")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")

    if missing:
        print("❌ Missing dependencies. Install them with:")
        print(f"   pip install {' '.join(missing)}")
        sys.exit(1)


def transcribe_audio(audio_path: str, model_size: str = "medium") -> dict:
    """
    Transcribe audio file using Whisper.
    Returns dict with 'text' and 'segments'.
    """
    import whisper

    print(f"\n🔊 Loading Whisper model: {model_size}")
    print("   (First run downloads the model ~1.5GB for 'medium' — please wait)\n")

    model = whisper.load_model(model_size)

    print(f"📂 Transcribing: {audio_path}")
    print("   This may take a few minutes depending on audio length...\n")

    start = time.time()
    result = model.transcribe(audio_path, verbose=False)
    elapsed = time.time() - start

    duration_mins = elapsed / 60
    print(f"✅ Transcription complete in {duration_mins:.1f} minutes\n")

    return result


def save_as_txt(result: dict, output_path: str):
    """Save plain text transcript."""
    with open(output_path, "w", encoding="utf-8") as f:
        # Write full transcript
        f.write("=" * 60 + "\n")
        f.write("TRANSCRIPT\n")
        f.write("=" * 60 + "\n\n")
        f.write(result["text"].strip())
        f.write("\n\n")

        # Write timestamped segments
        f.write("=" * 60 + "\n")
        f.write("TIMESTAMPED SEGMENTS\n")
        f.write("=" * 60 + "\n\n")

        for seg in result.get("segments", []):
            start = format_time(seg["start"])
            end = format_time(seg["end"])
            text = seg["text"].strip()
            f.write(f"[{start} --> {end}]  {text}\n")

    print(f"📄 TXT saved: {output_path}")


def save_as_docx(result: dict, output_path: str, source_file: str):
    """Save formatted Word document transcript."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title ──────────────────────────────────────────────
    title = doc.add_heading("Audio Transcript", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadata ───────────────────────────────────────────
    from datetime import datetime
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Source: {Path(source_file).name}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ── Full Transcript ────────────────────────────────────
    doc.add_heading("Full Transcript", level=2)
    para = doc.add_paragraph(result["text"].strip())
    para.style.font.size = Pt(11)

    doc.add_page_break()

    # ── Timestamped Segments ───────────────────────────────
    doc.add_heading("Timestamped Segments", level=2)

    for seg in result.get("segments", []):
        start = format_time(seg["start"])
        end = format_time(seg["end"])
        text = seg["text"].strip()

        p = doc.add_paragraph()

        # Timestamp in bold grey
        ts_run = p.add_run(f"[{start} → {end}]  ")
        ts_run.bold = True
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Segment text
        text_run = p.add_run(text)
        text_run.font.size = Pt(10)

        p.paragraph_format.space_after = Pt(4)

    doc.save(output_path)
    print(f"📝 DOCX saved: {output_path}")


def format_time(seconds: float) -> str:
    """Format seconds to MM:SS."""
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{mins:02d}:{secs:02d}"


def main():
    parser = argparse.ArgumentParser(
        description="Convert MP3 audio to transcript using Whisper"
    )
    parser.add_argument("audio", help="Path to the MP3 file")
    parser.add_argument(
        "--output", "-o",
        help="Output filename (without extension). Default: same as input file",
        default=None
    )
    parser.add_argument(
        "--format", "-f",
        choices=["txt", "docx", "both"],
        default="both",
        help="Output format (default: both)"
    )
    parser.add_argument(
        "--model", "-m",
        default="medium",
        choices=["tiny", "base", "small", "medium", "large"],
        help="Whisper model size (default: medium)"
    )

    args = parser.parse_args()

    # Validate input file
    audio_path = Path(args.audio)
    if not audio_path.exists():
        print(f"❌ File not found: {audio_path}")
        sys.exit(1)

    if audio_path.suffix.lower() not in [".mp3", ".wav", ".m4a", ".ogg", ".flac"]:
        print(f"⚠️  Warning: File extension '{audio_path.suffix}' may not be supported.")
        print("   Supported: .mp3, .wav, .m4a, .ogg, .flac")

    # Determine output base name
    output_base = args.output if args.output else audio_path.stem

    # Check deps
    check_dependencies()

    # Transcribe
    result = transcribe_audio(str(audio_path), model_size=args.model)

    # Save outputs
    print("💾 Saving output(s)...\n")

    if args.format in ("txt", "both"):
        save_as_txt(result, f"{output_base}.txt")

    if args.format in ("docx", "both"):
        save_as_docx(result, f"{output_base}.docx", str(audio_path))

    print("\n✅ Done!")


if __name__ == "__main__":
    main()
